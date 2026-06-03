"""
Telegram-бот — главный интерфейс Virtual CEO.
Планировщик: запрос балансов в 23:00, повтор в 12:00 если нет ответа.
Поддержка голосовых сообщений через OpenAI Whisper.
"""

import os
import json
import logging
import tempfile
from datetime import datetime, time, date
from io import BytesIO
from pathlib import Path

import pytz
from dotenv import load_dotenv
from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import (
    Application, CommandHandler, MessageHandler,
    CallbackQueryHandler, filters, ContextTypes
)

load_dotenv()

Path("logs").mkdir(exist_ok=True)
logging.basicConfig(
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
    level=logging.INFO,
    handlers=[
        logging.FileHandler("logs/bot.log"),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

OWNER_ID = int(os.getenv("TELEGRAM_OWNER_ID", "0"))
SALESPERSON_ID = int(os.getenv("SALESPERSON_ID", "0"))
ALMATY_TZ = pytz.FixedOffset(6 * 60)  # UTC+6 (фактический часовой пояс системы)
STATE_FILE = Path(__file__).parent.parent / "data" / "daily_state.json"

conversation_history: list = []
statement_analyses: list = []

ADDDATA_FILE = Path(__file__).parent.parent / "data" / "adddata_state.json"

kpiset_state: dict = {"active": False, "step": None, "temp": {}}  # in-memory
kpicalc_state: dict = {"active": False, "step": None, "temp": {}}  # in-memory
invoice_state: dict = {"active": False, "step": None, "temp": {}}  # in-memory

def _load_kpiset() -> dict:
    return kpiset_state

def _save_kpiset(updates: dict):
    global kpiset_state
    kpiset_state.update(updates)

def _reset_kpiset():
    global kpiset_state
    kpiset_state.update({"active": False, "step": None, "temp": {}})

def _load_kpicalc() -> dict:
    return kpicalc_state

def _save_kpicalc(updates: dict):
    global kpicalc_state
    kpicalc_state.update(updates)

def _reset_kpicalc():
    global kpicalc_state
    kpicalc_state.update({"active": False, "step": None, "temp": {}})

def _load_invoice() -> dict:
    return invoice_state

def _save_invoice(updates: dict):
    global invoice_state
    invoice_state.update(updates)

def _reset_invoice():
    global invoice_state
    invoice_state.update({"active": False, "step": None, "temp": {}})

YEARPLAN_FILE = Path(__file__).parent.parent / "data" / "yearplan_state.json"
ANALYST_STATE_FILE = Path(__file__).parent.parent / "data" / "analyst_state.json"
ANALYST_HISTORY_FILE = Path(__file__).parent.parent / "data" / "analyst_history.json"
STATEMENT_PENDING_FILE = Path(__file__).parent.parent / "data" / "statement_pending.json"
TOWORD_STATE_FILE = Path(__file__).parent.parent / "data" / "toword_state.json"
BRIEF_STATE_FILE = Path(__file__).parent.parent / "data" / "brief_state.json"

def _load_yearplan() -> dict:
    if YEARPLAN_FILE.exists():
        with open(YEARPLAN_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    return {"active": False, "data": [], "year": None, "waiting_year": False}

def _save_yearplan(state: dict):
    with open(YEARPLAN_FILE, "w", encoding="utf-8") as f:
        json.dump(state, f, ensure_ascii=False, indent=2)


# ─── Аналитик ─────────────────────────────────────────────────────────────────

def _analyst_active() -> bool:
    if not ANALYST_STATE_FILE.exists():
        return False
    with open(ANALYST_STATE_FILE, "r", encoding="utf-8") as f:
        return json.load(f).get("active", False)

def _analyst_set_active(value: bool):
    with open(ANALYST_STATE_FILE, "w", encoding="utf-8") as f:
        json.dump({"active": value}, f)

def _analyst_load_history() -> list:
    if not ANALYST_HISTORY_FILE.exists():
        return []
    with open(ANALYST_HISTORY_FILE, "r", encoding="utf-8") as f:
        return json.load(f)

def _analyst_save_history(history: list):
    with open(ANALYST_HISTORY_FILE, "w", encoding="utf-8") as f:
        json.dump(history, f, ensure_ascii=False, indent=2)

def _toword_active() -> bool:
    return TOWORD_STATE_FILE.exists()

def _toword_set(value: bool):
    if value:
        TOWORD_STATE_FILE.write_text("1")
    elif TOWORD_STATE_FILE.exists():
        TOWORD_STATE_FILE.unlink()

def _brief_load() -> dict:
    if not BRIEF_STATE_FILE.exists():
        return {}
    try:
        return json.loads(BRIEF_STATE_FILE.read_text(encoding="utf-8"))
    except Exception:
        return {}

def _brief_save(data: dict):
    BRIEF_STATE_FILE.write_text(json.dumps(data, ensure_ascii=False), encoding="utf-8")

def _brief_clear():
    if BRIEF_STATE_FILE.exists():
        BRIEF_STATE_FILE.unlink()

def _brief_active() -> bool:
    return bool(_brief_load().get("step"))

def _brief_set(value: bool):
    if not value:
        _brief_clear()

def _statement_pending_save(pdf_b64: str, bank: str):
    with open(STATEMENT_PENDING_FILE, "w", encoding="utf-8") as f:
        json.dump({"pdf_b64": pdf_b64, "bank": bank}, f)

def _statement_pending_load() -> dict:
    if not STATEMENT_PENDING_FILE.exists():
        return {}
    try:
        with open(STATEMENT_PENDING_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return {}

def _statement_pending_clear():
    if STATEMENT_PENDING_FILE.exists():
        STATEMENT_PENDING_FILE.unlink()

def _analyst_clear_history():
    with open(ANALYST_HISTORY_FILE, "w", encoding="utf-8") as f:
        json.dump([], f)


def _build_analyst_context() -> str:
    """Собирает актуальные бизнес-данные для системного промпта аналитика."""
    parts = []
    try:
        from agents.finance_agent import _load as _load_fin, format_summary
        fin = _load_fin()
        parts.append(f"=== ЛИЧНЫЕ СЧЕТА ОСНОВАТЕЛЯ ===\n{format_summary(fin)}")
    except Exception:
        pass
    try:
        from agents.sales_agent import get_historical_two_years
        parts.append(f"=== ВЫРУЧКА ПРОЕКТОВ (история) ===\n{get_historical_two_years()}")
    except Exception:
        pass
    try:
        kpi_file = Path(__file__).parent.parent / "data" / "kpi_plans.json"
        if kpi_file.exists():
            with open(kpi_file, "r", encoding="utf-8") as f:
                parts.append(f"=== KPI ПЛАНЫ ===\n{json.dumps(json.load(f), ensure_ascii=False, indent=2)}")
    except Exception:
        pass
    try:
        from agents.invoice_agent import _load_invoices
        inv = _load_invoices()
        invoices = inv.get("invoices", [])
        total = sum(float(i.get("amount", 0)) for i in invoices)
        parts.append(f"=== СЧЕТА НА ОПЛАТУ ===\nВсего: {len(invoices)}, сумма: {total:,.0f} ₸\nПоследние: {json.dumps(invoices[-5:], ensure_ascii=False)}")
    except Exception:
        pass
    return "\n\n".join(parts)


ANALYST_SYSTEM = """Ты бизнес-аналитик и советник Ибакдаулета — казахстанского медиа-предпринимателя, основателя Kettik Group.

Проекты:
- Grants KZ: гранты для казахстанцев, реклама в Instagram/Telegram, ~1.5М ₸/мес
- Tanda Bilim: история Казахстана (видео), реклама, ~1.5М ₸/мес
- Ekonomist Media: экономика/финансы, молодой проект, пока не зарабатывает
- Kettik Group: медиа-холдинг ~70 человек, ~75М ₸/мес (доля 35%)

Твоя роль: анализировать данные, находить точки роста, давать конкретные рекомендации с цифрами.
Помни контекст предыдущих сообщений в этом разговоре.
Отвечай на русском, кратко и по делу."""


# ─── Утилиты ────────────────────────────────────────────────────────────────

def is_owner(update: Update) -> bool:
    return update.effective_user.id == OWNER_ID


def is_salesperson(update: Update) -> bool:
    return SALESPERSON_ID != 0 and update.effective_user.id == SALESPERSON_ID


async def send_long(update: Update, text: str):
    MAX_LEN = 4000
    for chunk in [text[i:i+MAX_LEN] for i in range(0, len(text), MAX_LEN)]:
        await update.message.reply_text(chunk)


def load_state() -> dict:
    with open(STATE_FILE, "r") as f:
        return json.load(f)


def save_state(state: dict):
    with open(STATE_FILE, "w") as f:
        json.dump(state, f, ensure_ascii=False, indent=2)


def mark_responded():
    state = load_state()
    state["pending_update"] = False
    state["last_update"] = datetime.now(ALMATY_TZ).isoformat()
    save_state(state)


# ─── Продажи: запрос, напоминания, отчёт ────────────────────────────────────

async def request_sales_report(context):
    """17:00 — запускаем пошаговый диалог с продажником."""
    if SALESPERSON_ID == 0:
        return
    from agents.sales_agent import SalesConversation, reset_daily_state
    reset_daily_state()
    conv = SalesConversation()
    conv.reset()
    msg = conv.start()
    await context.bot.send_message(chat_id=SALESPERSON_ID, text=msg)


async def check_stale_conversation(context):
    """Каждые 30 мин: если диалог завис — сброс и напоминание продажнику."""
    if SALESPERSON_ID == 0:
        return
    from agents.sales_agent import SalesConversation, reset_daily_state
    conv = SalesConversation()
    if conv.is_stale(max_minutes=60):
        reset_daily_state()
        conv.reset()
        msg = conv.start()
        await context.bot.send_message(
            chat_id=SALESPERSON_ID,
            text="⚠️ Отчёт не был заполнен до конца. Начинаем заново — нужно заполнить все 3 проекта.\n\n" + msg
        )
        logger.info("Зависший диалог сброшен и перезапущен")


async def check_sales_19(context):
    """19:00 — если отчёт получен, шлём владельцу. Если нет — напоминание."""
    from agents.sales_agent import is_submitted_today, daily_report, get_reminder_message
    if is_submitted_today():
        report = daily_report()
        await context.bot.send_message(chat_id=OWNER_ID, text=report)
    else:
        if SALESPERSON_ID != 0:
            await context.bot.send_message(
                chat_id=SALESPERSON_ID,
                text=get_reminder_message(attempt=1)
            )
        await context.bot.send_message(
            chat_id=OWNER_ID,
            text="⏳ Дневной отчёт по продажам ещё не получен. Напомнил продажнику."
        )


async def remind_sales_2030(context):
    """20:30 — второе напоминание если всё ещё нет данных."""
    from agents.sales_agent import is_submitted_today, get_reminder_message
    if is_submitted_today():
        return
    if SALESPERSON_ID != 0:
        await context.bot.send_message(
            chat_id=SALESPERSON_ID,
            text=get_reminder_message(attempt=2)
        )
    await context.bot.send_message(
        chat_id=OWNER_ID,
        text="🔔 Продажник до сих пор не прислал отчёт."
    )


async def weekly_sales_job(context):
    """Понедельник 9:00 — недельный отчёт владельцу."""
    from agents.sales_agent import weekly_report
    report = weekly_report()
    await context.bot.send_message(chat_id=OWNER_ID, text=report)


async def monthly_sales_job(context):
    """1-е число 9:00 — месячный отчёт владельцу."""
    from agents.sales_agent import monthly_report
    import datetime as _dt
    last_month = date.today().replace(day=1) - _dt.timedelta(days=1)
    report = monthly_report(last_month.strftime("%Y-%m"))
    await context.bot.send_message(chat_id=OWNER_ID, text=report)


def _make_keyboard(buttons: list) -> InlineKeyboardMarkup:
    """Строит InlineKeyboardMarkup из списка [(текст, data), ...]."""
    keyboard = [
        [InlineKeyboardButton(text, callback_data=data) for text, data in row]
        for row in buttons
    ]
    return InlineKeyboardMarkup(keyboard)


async def handle_salesperson_message(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Пошаговый диалог с продажником."""
    from agents.sales_agent import SalesConversation
    conv = SalesConversation()
    text = update.message.text.strip().lower()

    # Продажник в режиме создания счёта
    inv = _load_invoice()
    if inv["active"]:
        await _handle_invoice_dialog(update)
        return

    if not conv.is_active():
        # Ждём "да" чтобы начать
        if text in ["да", "yes", "готов", "ок", "ok", "давай", "конечно"]:
            msg = conv.start()
            await update.message.reply_text(msg)
        else:
            await update.message.reply_text("Напиши 'да' когда будешь готов заполнить отчёт.")
        return

    await update.message.reply_chat_action("typing")
    next_q, is_done, summary = conv.process_answer(update.message.text.strip())
    await update.message.reply_text(next_q)

    if is_done and summary:
        await context.bot.send_message(chat_id=OWNER_ID, text=summary)


# ─── Курс доллара ────────────────────────────────────────────────────────────

async def auto_update_rate(context):
    """Ежедневно обновляет курс USD/KZT автоматически."""
    from agents.finance_agent import update_usd_rate
    result = update_usd_rate()
    logger.info(f"Курс: {result}")


# ─── Планировщик ────────────────────────────────────────────────────────────

async def request_balances(context):
    """Отправляет запрос балансов владельцу."""
    state = load_state()
    now = datetime.now(ALMATY_TZ)
    today_str = now.strftime("%Y-%m-%d")

    # Не беспокоим если уже обновили сегодня
    if today_str in state.get("last_update", ""):
        return

    from agents.finance_agent import format_summary, _load
    data = _load()
    summary = format_summary(data)

    state["pending_update"] = True
    state["last_request"] = now.isoformat()
    save_state(state)

    text = (
        f"🌙 Добрый вечер! Время обновить балансы.\n\n"
        f"Текущее состояние:\n{summary}\n\n"
        f"Отправьте новые балансы текстом или голосовым.\n"
        f"Пример: каспи 150000, халык 800000, халык депозит 24500000"
    )
    await context.bot.send_message(chat_id=OWNER_ID, text=text)


async def reminder_balances(context):
    """Повторный запрос в 12:00 если не ответил вчера вечером."""
    state = load_state()
    now = datetime.now(ALMATY_TZ)
    today_str = now.strftime("%Y-%m-%d")

    # Обновил сегодня — не беспокоим
    if today_str in state.get("last_update", ""):
        return
    # Нет pending — уже всё ок
    if not state.get("pending_update"):
        return

    await context.bot.send_message(
        chat_id=OWNER_ID,
        text=(
            "☀️ Напоминание: вчера вечером не обновили балансы счетов.\n"
            "Отправьте текстом или голосовым когда будет время."
        )
    )


# ─── Голосовые сообщения ─────────────────────────────────────────────────────

async def handle_voice(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    if not is_owner(update):
        return

    openai_key = os.getenv("OPENAI_API_KEY")
    if not openai_key:
        await update.message.reply_text(
            "Голосовые сообщения требуют OPENAI_API_KEY в .env файле."
        )
        return

    await update.message.reply_text("🎙 Слушаю...")

    tmp_ogg = tmp_mp3 = None
    try:
        from openai import OpenAI
        oai = OpenAI(api_key=openai_key)

        voice = update.message.voice
        file = await context.bot.get_file(voice.file_id)

        with tempfile.NamedTemporaryFile(suffix=".ogg", delete=False) as f:
            tmp_ogg = f.name
        await file.download_to_drive(tmp_ogg)

        # Транскрибируем через Whisper
        with open(tmp_ogg, "rb") as audio:
            transcript = oai.audio.transcriptions.create(
                model="whisper-1",
                file=audio,
                language="ru"
            )

        text = transcript.text
        await update.message.reply_text(f"🎙 Распознал: {text}\n\nОбновляю...")
        await _process_balance_text(update, text)

    except Exception as e:
        logger.error(f"Ошибка голосового: {e}")
        await update.message.reply_text(f"Не смог распознать: {e}")
    finally:
        for p in [tmp_ogg, tmp_mp3]:
            if p:
                try:
                    os.unlink(p)
                except Exception:
                    pass


async def _process_balance_text(update: Update, text: str):
    """Обновляет балансы из текста или голосовой транскрипции через Claude."""
    import anthropic as ant
    import json as _json
    from agents.finance_agent import update_balance, _load, format_summary

    data = _load()

    # Claude извлекает балансы — понимает и цифры и слова ("сто пятьдесят тысяч")
    client = ant.Anthropic(api_key=os.getenv("ANTHROPIC_API_KEY"))

    accounts_list = "\n".join(
        f'  "{k}": "{v["name"]} ({v["currency"]})"'
        for k, v in data["accounts"].items()
    )

    prompt = f"""Из этого текста извлеки балансы счетов. Текст может быть голосовым (числа словами).

Текст: "{text}"

Доступные счета:
{accounts_list}

Верни JSON массив только тех счетов которые упомянуты:
[{{"account_key": "kaspi_card", "balance": 150000}}, ...]

Правила:
- "сто пятьдесят тысяч" = 150000
- "миллион пятьсот" = 1500000
- Если счёт не упомянут — не включай
- Только JSON, без текста"""

    resp = client.messages.create(
        model="claude-opus-4-6",
        max_tokens=512,
        messages=[{"role": "user", "content": prompt}]
    )

    raw = resp.content[0].text.strip()

    # Извлекаем JSON из ответа
    import re
    match = re.search(r'\[.*\]', raw, re.DOTALL)
    if not match:
        from core.orchestrator import run
        response = run(text, conversation_history)
        await send_long(update, response)
        return

    try:
        balances = _json.loads(match.group())
    except Exception:
        await update.message.reply_text("Не смог распознать балансы. Попробуйте ещё раз.")
        return

    if not balances:
        from core.orchestrator import run
        response = run(text, conversation_history)
        await send_long(update, response)
        return

    updated = []
    for item in balances:
        key = item.get("account_key")
        bal = item.get("balance")
        if key and bal is not None:
            result = update_balance(key, float(bal), data)
            updated.append(result)

    if updated:
        mark_responded()
        summary = format_summary(data)
        await update.message.reply_text(
            "✅ Балансы обновлены:\n" + "\n".join(updated) + f"\n\n{summary}"
        )
    else:
        await update.message.reply_text("Не нашёл балансов в сообщении. Попробуйте ещё раз.")


# ─── Команды ─────────────────────────────────────────────────────────────────

async def start(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    if not is_owner(update):
        return
    await update.message.reply_text(
        "Виртуальный СЕО запущен. v2.1\n\n"
        "Команды:\n"
        "/clear — очистить диалог\n"
        "/reset — сбросить загруженные выписки\n"
        "/balances — запросить балансы сейчас\n"
        "/adddata 2026 — добавить/изменить данные продаж\n"
        "/bonus — рассчитать бонусы продажника\n"
        "/plan — установить KPI план на месяц\n"
        "/sales — запросить отчёт у продажника\n"
        "/yearplan — составить годовой план"
    )


async def cmd_clear(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    if not is_owner(update):
        return
    conversation_history.clear()
    await update.message.reply_text("История диалога очищена.")


async def cmd_reset(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    if not is_owner(update):
        return
    statement_analyses.clear()
    await update.message.reply_text("Выписки сброшены.")


async def cmd_balances(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    """Ручной запрос обновления балансов."""
    if not is_owner(update):
        return
    await request_balances(context)


def _load_adddata() -> dict:
    if ADDDATA_FILE.exists():
        with open(ADDDATA_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    return {"active": False, "year": None, "month": None, "step": None, "temp": {}}

def _save_adddata(updates: dict):
    state = _load_adddata()
    state.update(updates)
    with open(ADDDATA_FILE, "w", encoding="utf-8") as f:
        json.dump(state, f, ensure_ascii=False, indent=2)

def _reset_adddata():
    with open(ADDDATA_FILE, "w", encoding="utf-8") as f:
        json.dump({"active": False, "year": None, "month": None, "step": None, "temp": {}}, f)


async def cmd_kpi(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    """Пошаговый расчёт KPI — вводишь выручку сам."""
    if not is_owner(update):
        return
    ym = date.today().strftime("%Y-%m")
    _save_kpicalc({"active": True, "step": "ask_month", "temp": {}})
    await update.message.reply_text(
        f"💰 Расчёт бонусов\n\n"
        f"За какой месяц? Напиши период или 'текущий':\n"
        f"Пример: 2026-04 или текущий (сейчас: {ym})"
    )


async def cmd_setkpi(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    """Пошаговая установка KPI плана."""
    if not is_owner(update):
        return
    _save_kpiset({"active": True, "step": "month", "temp": {}})
    await update.message.reply_text(
        "📊 Установка KPI плана\n\n"
        "За какой месяц? Напиши год-месяц:\n"
        "Пример: 2026-05"
    )


async def cmd_adddata(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    """Пошаговое добавление/корректировка исторических продаж."""
    if not is_owner(update):
        return
    args = context.args
    if args and args[0].isdigit() and len(args[0]) == 4:
        year = int(args[0])
        _save_adddata({"active": True, "year": year, "month": None, "step": "ask_month", "temp": {}})
        from agents.sales_agent import get_historical_raw
        current = get_historical_raw(year)
        await update.message.reply_text(
            f"{current}\n\n"
            f"─────────────────\n"
            f"Какой месяц добавляем или корректируем?\n"
            f"Напиши название: Май, Июнь, Январь...\n"
            f"Или 'отмена' чтобы выйти."
        )
    else:
        await update.message.reply_text("Укажи год: /adddata 2025 или /adddata 2026")


async def cmd_yearplan(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    """Запускает режим составления годового плана продаж."""
    if not is_owner(update):
        return
    state = {"active": True, "data": [], "year": None, "waiting_year": True}
    _save_yearplan(state)
    await update.message.reply_text(
        "📊 Годовой план продаж.\n\nЗа какой год составляем план? (напиши год, например: 2026)"
    )


async def cmd_syncsheets(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    """Ручная синхронизация исторических данных в Google Sheets."""
    if not is_owner(update):
        return
    await update.message.reply_text("⏳ Синхронизирую таблицы...")
    try:
        from agents.sheets_agent import sync_historical_to_sheets
        sync_historical_to_sheets()
        await update.message.reply_text("✅ Google Sheets обновлён — листы История 2025 и История 2026.")
    except Exception as e:
        await update.message.reply_text(f"❌ Ошибка: {e}")


async def _extract_invoice_data(text: str) -> dict | None:
    """Claude извлекает данные счёта из свободного текста."""
    import anthropic as _ant, json as _json, re as _re
    client = _ant.Anthropic(api_key=os.getenv("ANTHROPIC_API_KEY"))
    prompt = f"""Из текста извлеки данные для счёта на оплату.

Текст: "{text}"

Верни JSON:
{{
  "invoice_number": 35,
  "client_name": "ТОО Astana Publicity",
  "client_bin": "100540014078",
  "client_address": "г. Астана, ул. Туркестан 28А",
  "service_name": "Реклама в Grants.kz",
  "service_code": "00000000150",
  "amount": 150000
}}

Сервис-коды по умолчанию:
- Реклама в Grants.kz → 00000000150
- Реклама в Tanda Bilim → 00000000151
- Реклама в Ekonomist Media → 00000000152
- Другое → 00000000000

Если в тексте явно указан код услуги (например "код 68") — используй его, дополни нулями до 11 знаков (68 → 00000000068).

Правила:
- invoice_number: номер счёта из текста (например "счёт 35", "№35", "номер 35") — если не указан, верни null
- "150к" = 150000, "1.5М" = 1500000
- Если client_name, client_bin, client_address, service_name или amount не найдены — верни null
- Только JSON"""
    resp = client.messages.create(
        model="claude-opus-4-6", max_tokens=512,
        messages=[{"role": "user", "content": prompt}]
    )
    raw = resp.content[0].text.strip()
    if raw.lower() == "null":
        return None
    m = _re.search(r'\{.*\}', raw, _re.DOTALL)
    if not m:
        return None
    try:
        return _json.loads(m.group())
    except Exception:
        return None


async def _handle_invoice_dialog(update: Update):
    """Обрабатывает сообщения пользователя в режиме создания счёта."""
    user_text = update.message.text.strip()
    inv = _load_invoice()

    if user_text.lower() in ["отмена", "стоп", "cancel"]:
        _reset_invoice()
        await update.message.reply_text("Отменено.")
        return

    step = inv["step"]
    if step == "waiting_text":
        await update.message.reply_chat_action("typing")
        data = await _extract_invoice_data(user_text)
        required = ["invoice_number", "client_name", "client_bin", "client_address", "service_name", "amount"]
        if data and all(data.get(k) for k in required):
            _reset_invoice()
            await _generate_and_send_invoice(update, data)
        else:
            missing = []
            if not data or not data.get("invoice_number"): missing.append("номер счёта")
            if not data or not data.get("client_name"): missing.append("название клиента")
            if not data or not data.get("client_bin"): missing.append("БИН")
            if not data or not data.get("client_address"): missing.append("адрес")
            if not data or not data.get("service_name"): missing.append("услуга")
            if not data or not data.get("amount"): missing.append("сумма")
            await update.message.reply_text(
                f"Не хватает данных: {', '.join(missing)}\n\n"
                f"Напиши ещё раз с полными данными или 'отмена'."
            )


async def _fetch_post(url: str, session) -> str:
    """Извлекает текст одной публикации (Instagram пост, Telegram пост, страница)."""
    import re
    url = url.strip().rstrip("/")
    headers = {"User-Agent": "Mozilla/5.0 (compatible; TelegramBot/1.0)"}

    try:
        # Instagram пост — oEmbed API (caption без авторизации)
        if "instagram.com/p/" in url or "instagram.com/reel/" in url:
            oembed_url = f"https://api.instagram.com/oembed/?url={url}"
            async with session.get(oembed_url, timeout=aiohttp.ClientTimeout(total=5)) as resp:
                if resp.status == 200:
                    data = await resp.json()
                    return data.get("title", "") or data.get("author_name", "")
            # Fallback: meta description
            async with session.get(url, headers=headers, timeout=aiohttp.ClientTimeout(total=5)) as resp:
                html = await resp.text(errors="ignore")
            m = re.search(r'<meta[^>]+property="og:description"[^>]+content="([^"]+)"', html)
            return m.group(1) if m else "(пост недоступен)"

        # Telegram пост — preview страница
        tg = re.match(r"(?:https?://)?t\.me/([^/]+)/(\d+)", url)
        if tg:
            channel, msg_id = tg.group(1), tg.group(2)
            embed_url = f"https://t.me/{channel}/{msg_id}?embed=1"
            async with session.get(embed_url, headers=headers, timeout=aiohttp.ClientTimeout(total=5)) as resp:
                html = await resp.text(errors="ignore")
            found = re.findall(r'class="tgme_widget_message_text[^"]*"[^>]*>(.*?)</div>', html, re.DOTALL)
            if found:
                return re.sub(r"<[^>]+>", "", found[0]).strip()
            return "(пост недоступен)"

        # Любая другая страница — og:description или первый абзац
        async with session.get(url, headers=headers, timeout=aiohttp.ClientTimeout(total=6)) as resp:
            html = await resp.text(errors="ignore")
        m = re.search(r'<meta[^>]+property="og:description"[^>]+content="([^"]+)"', html)
        if m:
            return m.group(1)
        m = re.search(r'<p[^>]*>([^<]{50,})</p>', html)
        return m.group(1).strip() if m else "(текст не найден)"

    except Exception as e:
        return f"(ошибка: {e})"


async def _scrape_post_urls(urls: list) -> str:
    """Параллельно читает несколько ссылок на публикации и возвращает их тексты."""
    import aiohttp, asyncio
    async with aiohttp.ClientSession() as session:
        tasks = [_fetch_post(url, session) for url in urls[:5]]
        results = await asyncio.gather(*tasks)
    parts = []
    for url, text in zip(urls, results):
        if text and "(ошибка" not in text and "(пост недоступен)" not in text:
            parts.append(f"📌 {url}\n{text}")
    return "\n\n---\n\n".join(parts) if parts else "Не удалось извлечь тексты публикаций."


async def _analyze_client_profile(url: str, page_text: str) -> str:
    """Анализирует контент страницы и составляет портрет клиента."""
    import anthropic as _ant, asyncio
    def _call():
        client = _ant.Anthropic(api_key=os.getenv("ANTHROPIC_API_KEY"))
        resp = client.messages.create(
            model="claude-opus-4-6",
            max_tokens=1000,
            messages=[{"role": "user", "content": f"""Проанализируй страницу клиента ({url}):

{page_text}

На основе последних постов составь портрет клиента:
- Ниша и продукт/услуга
- Tone of voice (формальный/дружелюбный/экспертный/юморной)
- Стиль текстов (длина, структура, эмодзи, обращение)
- Целевая аудитория (кто их читатели)
- Ключевые темы и ценности бренда
- Что они НЕ публикуют (чего избегать в рекламе)

Максимум 250 слов."""}]
        )
        return resp.content[0].text
    return await asyncio.to_thread(_call)


async def _generate_ad_from_brief(brief_text: str, client_profile: str = "", client_text: str = "") -> str:
    """Генерирует рекламный контент на основе брифа + портрета клиента + текста."""
    import anthropic as _ant, asyncio

    context_parts = []
    if client_profile:
        context_parts.append(f"КОНТЕНТ СО СТРАНИЦЫ КЛИЕНТА (изучи стиль, тон, аудиторию):\n{client_profile}")
    if client_text:
        context_parts.append(f"ГОТОВЫЙ ТЕКСТ ОТ КЛИЕНТА:\n{client_text}")
    context_parts.append(f"ЗАПОЛНЕННЫЙ БРИФ:\n{brief_text}")
    full_context = "\n\n---\n\n".join(context_parts)

    prompt = f"""Ты — опытный SMM-копирайтер медиапроекта Grants KZ (зарубежные гранты и стипендии для казахстанцев, аудитория 18-30 лет).

{full_context}

На основе всей информации создай рекламный пакет:

---
📸 INSTAGRAM ПОСТ
(эмодзи, живой язык, обращение на «ты», хук в первых 2 строках, хэштеги в конце)

---
✈️ TELEGRAM ПОСТ
(без хэштегов, чуть более информационный, но живой, эмодзи уместно)

---
🎯 3 ВАРИАНТА ХУКА
(цепляющие первые строки — выбери лучший)

---
📣 CTA БЛОК
(призыв к действию согласно брифу)

---

Учти стиль клиента, тон брифа и целевую аудиторию. Пиши на русском языке."""

    def _call():
        client = _ant.Anthropic(api_key=os.getenv("ANTHROPIC_API_KEY"))
        resp = client.messages.create(
            model="claude-opus-4-6",
            max_tokens=3000,
            messages=[{"role": "user", "content": prompt}]
        )
        return resp.content[0].text
    return await asyncio.to_thread(_call)


async def cmd_brief(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    """Активирует режим генерации рекламы по брифу."""
    if not is_owner(update):
        return
    args = context.args
    if args and args[0].lower() == "stop":
        _brief_clear()
        await update.message.reply_text("Режим брифа выключен.")
        return
    _brief_save({"step": "waiting_url", "client_profile": "", "client_text": ""})
    await update.message.reply_text(
        "📋 Создаём рекламу по брифу.\n\n"
        "Шаг 1/3: Отправь 2-5 ссылок на публикации клиента (каждая на новой строке) — изучу стиль и тон их контента.\n\n"
        "Поддерживаются: Instagram посты, Telegram посты, страницы сайта.\n\n"
        "Если ссылок нет — напиши *пропустить*.",
        parse_mode="Markdown"
    )


async def _process_brief_file(update: Update, brief_content_blocks: list) -> None:
    """Извлекает текст брифа и генерирует рекламный контент."""
    import anthropic as _ant, asyncio
    state = _brief_load()
    _brief_clear()
    await update.message.reply_text("📋 Читаю бриф...")

    def _extract():
        client = _ant.Anthropic(api_key=os.getenv("ANTHROPIC_API_KEY"))
        resp = client.messages.create(
            model="claude-opus-4-6",
            max_tokens=2000,
            messages=[{"role": "user", "content": brief_content_blocks + [
                {"type": "text", "text": "Извлеки все заполненные поля брифа в виде структурированного текста. Формат: Поле: Значение. Пропусти пустые поля."}
            ]}]
        )
        return resp.content[0].text

    try:
        brief_text = await asyncio.to_thread(_extract)
        await update.message.reply_text("✍️ Генерирую рекламные тексты...")
        ad_content = await _generate_ad_from_brief(
            brief_text,
            client_profile=state.get("client_profile", ""),
            client_text=state.get("client_text", "")
        )
        await send_long(update, ad_content)
    except Exception as e:
        logger.error(f"[Brief] error: {e}")
        await update.message.reply_text(f"❌ Ошибка: {e}")


async def _make_word_from_text(text: str, filename: str = "document.docx") -> BytesIO:
    """Создаёт Word файл из текста."""
    from docx import Document
    from docx.shared import Pt
    doc = Document()
    for line in text.split("\n"):
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        else:
            doc.add_paragraph(line)
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


async def _extract_text_from_pdf_for_word(pdf_bytes: bytes) -> str:
    """Извлекает текст из PDF через Claude."""
    import base64, anthropic as _ant, asyncio
    pdf_b64 = base64.standard_b64encode(pdf_bytes).decode()
    def _call():
        client = _ant.Anthropic(api_key=os.getenv("ANTHROPIC_API_KEY"))
        resp = client.messages.create(
            model="claude-opus-4-6",
            max_tokens=4096,
            messages=[{"role": "user", "content": [
                {"type": "document", "source": {"type": "base64", "media_type": "application/pdf", "data": pdf_b64}},
                {"type": "text", "text": "Извлеки весь текст из этого документа. Сохрани структуру: заголовки обозначь через # ## ###, абзацы — пустой строкой. Только текст, без комментариев."}
            ]}]
        )
        return resp.content[0].text
    return await asyncio.to_thread(_call)


async def _extract_text_from_image_for_word(image_bytes: bytes, mime: str = "image/jpeg") -> str:
    """Извлекает текст из изображения через Claude vision."""
    import base64, anthropic as _ant, asyncio
    img_b64 = base64.standard_b64encode(image_bytes).decode()
    def _call():
        client = _ant.Anthropic(api_key=os.getenv("ANTHROPIC_API_KEY"))
        resp = client.messages.create(
            model="claude-opus-4-6",
            max_tokens=4096,
            messages=[{"role": "user", "content": [
                {"type": "image", "source": {"type": "base64", "media_type": mime, "data": img_b64}},
                {"type": "text", "text": "Извлеки весь текст с этого изображения. Сохрани структуру: заголовки обозначь через # ## ###, абзацы — пустой строкой. Только текст, без комментариев."}
            ]}]
        )
        return resp.content[0].text
    return await asyncio.to_thread(_call)


async def cmd_toword(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    """Включает режим конвертации PDF/фото → Word."""
    if not is_owner(update):
        return
    _toword_set(True)
    await update.message.reply_text(
        "📄 Режим конвертации включён.\n\n"
        "Отправь PDF или фото документа — верну готовый Word файл.\n"
        "/toword stop — выйти."
    )


async def handle_photo(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    """Обрабатывает фото — конвертирует в Word если режим активен."""
    if not is_owner(update):
        return
    # Режим брифа — фото брифа
    if _brief_load().get("step") == "waiting_brief":
        try:
            import base64
            photo = update.message.photo[-1]
            file = await context.bot.get_file(photo.file_id)
            with tempfile.NamedTemporaryFile(suffix=".jpg", delete=False) as tmp:
                tmp_path = tmp.name
            await file.download_to_drive(tmp_path)
            with open(tmp_path, "rb") as f:
                img_b64 = base64.standard_b64encode(f.read()).decode()
            os.unlink(tmp_path)
            blocks = [{"type": "image", "source": {"type": "base64", "media_type": "image/jpeg", "data": img_b64}}]
            await _process_brief_file(update, blocks)
        except Exception as e:
            await update.message.reply_text(f"❌ Ошибка: {e}")
        return

    if not _toword_active():
        return
    _toword_set(False)
    await update.message.reply_chat_action("upload_document")
    try:
        photo = update.message.photo[-1]
        file = await context.bot.get_file(photo.file_id)
        with tempfile.NamedTemporaryFile(suffix=".jpg", delete=False) as tmp:
            tmp_path = tmp.name
        await file.download_to_drive(tmp_path)
        with open(tmp_path, "rb") as f:
            image_bytes = f.read()
        os.unlink(tmp_path)
        await update.message.reply_text("🔍 Читаю текст с фото...")
        text = await _extract_text_from_image_for_word(image_bytes)
        fname = "document.docx"
        word_buf = await _make_word_from_text(text, fname)
        await update.message.reply_document(document=word_buf, filename=fname, caption="✅ Word файл готов")
    except Exception as e:
        logger.error(f"[toword photo] {e}")
        await update.message.reply_text(f"❌ Ошибка: {e}")


async def cmd_invoice(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    """Создание счёта — с текстом сразу или пошагово."""
    if not (is_owner(update) or is_salesperson(update)):
        return

    # Если передан текст сразу: /invoice ТОО Рога и Копыта, БИН 123...
    text = " ".join(context.args) if context.args else ""
    if text.strip():
        await update.message.reply_chat_action("typing")
        data = await _extract_invoice_data(text)
        if data and all(data.get(k) for k in ["invoice_number", "client_name", "client_bin", "client_address", "service_name", "amount"]):
            await _generate_and_send_invoice(update, data)
        else:
            await update.message.reply_text(
                "Не смог извлечь все данные. Укажи:\n"
                "клиент, БИН, адрес, услуга, сумма\n\n"
                "Или начни пошагово — напиши /invoice без текста."
            )
        return

    _save_invoice({"active": True, "step": "waiting_text", "temp": {}})
    await update.message.reply_text(
        "🧾 Счёт на оплату\n\n"
        "Напиши данные клиента — я сам всё извлеку:\n\n"
        "Пример:\n"
        "ТОО Astana Publicity, БИН 100540014078, г. Астана ул. Туркестан 28А, реклама Grants KZ, 150 000 тенге\n\n"
        "Если услуга нестандартная — добавь код:\n"
        "пост в телеграм, код 68, 80 000 тенге"
    )


async def _generate_and_send_invoice(update: Update, data: dict):
    """Генерирует PDF и отправляет в чат."""
    from io import BytesIO
    await update.message.reply_chat_action("upload_document")
    try:
        from agents.invoice_agent import generate_invoice_pdf, save_invoice_record
        num = data.get("invoice_number") or data.get("number")
        today = date.today()
        pdf_bytes = generate_invoice_pdf(
            invoice_number=num,
            invoice_date=today,
            client_name=data["client_name"],
            client_bin=data["client_bin"],
            client_address=data["client_address"],
            service_name=data["service_name"],
            service_code=data.get("service_code", "00000000000"),
            amount=float(data["amount"]),
        )
        save_invoice_record({
            "number": num,
            "date": today.isoformat(),
            "client": data["client_name"],
            "amount": data["amount"],
            "service": data["service_name"],
        })
        fname = f"Счет_{num}_{data['client_name'].replace(' ', '_')}.pdf"
        await update.message.reply_document(
            document=BytesIO(pdf_bytes),
            filename=fname,
            caption=f"🧾 Счёт №{num} | {data['client_name']} | {float(data['amount']):,.0f} ₸"
        )
    except Exception as e:
        logger.error(f"Invoice generation error: {e}")
        await update.message.reply_text(f"❌ Ошибка генерации счёта: {e}")


async def _handle_analyst_message(update: Update):
    """Обрабатывает сообщение в режиме аналитика."""
    import asyncio
    import anthropic as _ant

    user_text = update.message.text.strip()
    history = _analyst_load_history()
    history.append({"role": "user", "content": user_text})

    await update.message.reply_chat_action("typing")

    def _call():
        client = _ant.Anthropic(api_key=os.getenv("ANTHROPIC_API_KEY"))
        biz_context = _build_analyst_context()
        system = ANALYST_SYSTEM + f"\n\n--- АКТУАЛЬНЫЕ ДАННЫЕ ---\n{biz_context}"
        resp = client.messages.create(
            model="claude-opus-4-6",
            max_tokens=2000,
            system=system,
            messages=history[-20:],  # последние 20 сообщений
        )
        return resp.content[0].text

    try:
        answer = await asyncio.to_thread(_call)
        history.append({"role": "assistant", "content": answer})
        _analyst_save_history(history)
        await send_long(update, answer)
    except Exception as e:
        logger.error(f"[Analyst] error: {e}")
        await update.message.reply_text(f"❌ Ошибка аналитика: {e}")


async def cmd_analyst(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    """Управление режимом бизнес-аналитика."""
    if not is_owner(update):
        return

    arg = (context.args[0].lower() if context.args else "").strip()

    if arg == "stop":
        _analyst_set_active(False)
        await update.message.reply_text("🔴 Режим аналитика выключен.")
        return

    if arg == "reset":
        _analyst_clear_history()
        await update.message.reply_text("🗑 История аналитика очищена. Продолжай — я помню контекст с нуля.")
        return

    if arg == "history":
        history = _analyst_load_history()
        if not history:
            await update.message.reply_text("История пуста.")
            return
        lines = []
        for msg in history[-10:]:
            role = "Ты" if msg["role"] == "user" else "Аналитик"
            text = msg["content"][:200] + ("..." if len(msg["content"]) > 200 else "")
            lines.append(f"[{role}]: {text}")
        await send_long(update, "\n\n".join(lines))
        return

    # Включаем режим
    _analyst_set_active(True)
    history = _analyst_load_history()
    if history:
        await update.message.reply_text(
            f"🟢 Аналитик включён. Продолжаю прошлый разговор ({len(history)} сообщений).\n"
            f"Задай вопрос. /analyst stop — выйти, /analyst reset — очистить историю."
        )
    else:
        await update.message.reply_chat_action("typing")
        # Первый запуск — даём вводный анализ
        history.append({"role": "user", "content": "Привет! Дай краткий обзор текущего состояния бизнеса на основе имеющихся данных."})
        import asyncio, anthropic as _ant

        def _intro():
            client = _ant.Anthropic(api_key=os.getenv("ANTHROPIC_API_KEY"))
            biz_context = _build_analyst_context()
            system = ANALYST_SYSTEM + f"\n\n--- АКТУАЛЬНЫЕ ДАННЫЕ ---\n{biz_context}"
            resp = client.messages.create(
                model="claude-opus-4-6", max_tokens=2000,
                system=system, messages=history[-20:],
            )
            return resp.content[0].text

        try:
            answer = await asyncio.to_thread(_intro)
            history.append({"role": "assistant", "content": answer})
            _analyst_save_history(history)
            await send_long(update, f"🟢 Аналитик включён.\n\n{answer}\n\n─────\nЗадавай вопросы. /analyst stop — выйти.")
        except Exception as e:
            await update.message.reply_text(f"❌ Ошибка: {e}")


async def cmd_ceo(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    """Запускает CEO-анализ с выбором направления."""
    if not is_owner(update):
        return
    keyboard = _make_keyboard([
        [("💰 Финансы", "ceo:finance"), ("📈 Продажи", "ceo:sales")],
        [("🧾 Счета", "ceo:invoices"), ("🔭 Полная картина", "ceo:all")],
    ])
    await update.message.reply_text(
        "🤖 Virtual CEO — анализ\n\nВыбери направление:",
        reply_markup=keyboard,
    )


async def _ceo_analyze(topic: str, context_data: str) -> str:
    """Отправляет данные в Claude и получает стратегический анализ."""
    import asyncio
    import anthropic as _ant

    def _call():
        client = _ant.Anthropic(api_key=os.getenv("ANTHROPIC_API_KEY"))
        system = (
            "Ты виртуальный CEO медиа-холдинга Kettik Group (Казахстан). "
            "Анализируй данные кратко и по делу. Давай 3-5 конкретных советов "
            "с цифрами и приоритетами. Пиши на русском, без воды."
        )
        resp = client.messages.create(
            model="claude-opus-4-6",
            max_tokens=1500,
            system=system,
            messages=[{"role": "user", "content": f"Проанализируй данные и дай рекомендации.\n\nТема: {topic}\n\nДанные:\n{context_data}"}],
        )
        return resp.content[0].text

    return await asyncio.to_thread(_call)


async def _ceo_finance(update: Update):
    await update.message.reply_chat_action("typing")
    try:
        from agents.finance_agent import _load, format_summary
        data = _load()
        summary = format_summary(data)
        accounts_detail = json.dumps(data.get("accounts", {}), ensure_ascii=False, indent=2)
        context_data = (
            f"Это ЛИЧНЫЕ счета основателя (не компании).\n\n"
            f"Текущие балансы:\n{summary}\n\n"
            f"Детали:\n{accounts_detail}"
        )
        result = await _ceo_analyze(
            "Личные финансы основателя — текущие балансы личных счетов (Каспи, Халык и др.)",
            context_data,
        )
        await send_long(update, f"💰 Личные финансы\n{'─'*30}\n\n{result}")
    except Exception as e:
        await update.message.reply_text(f"❌ Ошибка: {e}")


async def _ceo_sales(update: Update):
    await update.message.reply_chat_action("typing")
    try:
        from agents.sales_agent import get_historical_two_years, daily_report, _load as _load_sales
        history = get_historical_two_years()
        sales_data = _load_sales()
        today_block = ""
        try:
            today_block = f"\nДневной отчёт (последний):\n{daily_report()}"
        except Exception:
            pass
        kpi_file = Path(__file__).parent.parent / "data" / "kpi_plans.json"
        kpi_block = ""
        if kpi_file.exists():
            with open(kpi_file, "r", encoding="utf-8") as f:
                kpi_block = f"\nKPI планы:\n{json.dumps(json.load(f), ensure_ascii=False, indent=2)}"
        context_data = f"{history}{today_block}{kpi_block}"
        result = await _ceo_analyze("Продажи — выручка, динамика, KPI", context_data)
        await send_long(update, f"📈 Анализ продаж\n{'─'*30}\n\n{result}")
    except Exception as e:
        await update.message.reply_text(f"❌ Ошибка: {e}")


async def _ceo_invoices(update: Update):
    await update.message.reply_chat_action("typing")
    try:
        from agents.invoice_agent import _load_invoices
        inv_data = _load_invoices()
        invoices = inv_data.get("invoices", [])
        if not invoices:
            await update.message.reply_text("📋 Счетов пока нет.")
            return
        total = sum(float(i.get("amount", 0)) for i in invoices)
        context_data = (
            f"Всего счетов: {len(invoices)}\n"
            f"Общая сумма: {total:,.0f} ₸\n\n"
            f"История:\n{json.dumps(invoices, ensure_ascii=False, indent=2)}"
        )
        result = await _ceo_analyze("Счета на оплату — дебиторская задолженность, клиенты", context_data)
        await send_long(update, f"🧾 Анализ счетов\n{'─'*30}\n\n{result}")
    except Exception as e:
        await update.message.reply_text(f"❌ Ошибка: {e}")


async def _ceo_all(update: Update):
    await update.message.reply_chat_action("typing")
    try:
        from agents.finance_agent import _load as _load_fin, format_summary
        from agents.sales_agent import get_historical_two_years, daily_report
        from agents.invoice_agent import _load_invoices

        fin_data = _load_fin()
        fin_summary = format_summary(fin_data)

        sales_history = get_historical_two_years()

        today_sales = ""
        try:
            today_sales = daily_report()
        except Exception:
            pass

        inv_data = _load_invoices()
        invoices = inv_data.get("invoices", [])
        inv_total = sum(float(i.get("amount", 0)) for i in invoices)

        kpi_file = Path(__file__).parent.parent / "data" / "kpi_plans.json"
        kpi_block = ""
        if kpi_file.exists():
            with open(kpi_file, "r", encoding="utf-8") as f:
                kpi_block = json.dumps(json.load(f), ensure_ascii=False, indent=2)

        context_data = (
            f"=== ФИНАНСЫ ===\n{fin_summary}\n\n"
            f"=== ПРОДАЖИ (история) ===\n{sales_history}\n\n"
            f"=== ПРОДАЖИ (сегодня) ===\n{today_sales}\n\n"
            f"=== СЧЕТА ===\nВсего: {len(invoices)}, сумма: {inv_total:,.0f} ₸\n\n"
            f"=== KPI ПЛАНЫ ===\n{kpi_block}"
        )
        result = await _ceo_analyze("Полная картина бизнеса — финансы, продажи, счета, KPI", context_data)
        await send_long(update, f"🔭 Полный анализ\n{'─'*30}\n\n{result}")
    except Exception as e:
        await update.message.reply_text(f"❌ Ошибка: {e}")


async def handle_ceo_callback(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    """Обрабатывает нажатия кнопок CEO-анализа."""
    query = update.callback_query
    await query.answer()
    if not query.data.startswith("ceo:"):
        return
    direction = query.data.split(":", 1)[1]
    labels = {"finance": "💰 Финансы", "sales": "📈 Продажи", "invoices": "🧾 Счета", "all": "🔭 Полная картина"}
    await query.edit_message_text(f"{labels.get(direction, direction)} — анализирую...")
    # Создаём фиктивный update чтобы reply_text работал корректно
    class _FakeUpdate:
        def __init__(self, q):
            self.message = q.message
            self.effective_user = q.from_user
    fake = _FakeUpdate(query)
    if direction == "finance":
        await _ceo_finance(fake)
    elif direction == "sales":
        await _ceo_sales(fake)
    elif direction == "invoices":
        await _ceo_invoices(fake)
    elif direction == "all":
        await _ceo_all(fake)


async def cmd_sales(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    """Внеплановый запрос дневного отчёта у продажника."""
    if not is_owner(update):
        return
    if SALESPERSON_ID == 0:
        await update.message.reply_text("SALESPERSON_ID не задан в .env")
        return
    from agents.sales_agent import SalesConversation, reset_daily_state
    reset_daily_state()
    conv = SalesConversation()
    conv.reset()
    msg = conv.start()
    await context.bot.send_message(chat_id=SALESPERSON_ID, text=msg)
    await update.message.reply_text("✅ Запрос отправлен продажнику.")


# ─── Текстовые сообщения ─────────────────────────────────────────────────────

async def handle_message(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    if is_salesperson(update):
        await handle_salesperson_message(update, context)
        return

    if not is_owner(update):
        return

    user_text = update.message.text.strip()
    await update.message.reply_chat_action("typing")

    # Многошаговый режим брифа
    brief_state = _brief_load()
    if brief_state.get("step") == "waiting_url":
        skip = user_text.lower() in ["пропустить", "пропусти", "skip", "-"]
        if skip:
            brief_state["client_profile"] = ""
        else:
            urls = [u.strip() for u in user_text.strip().splitlines() if u.strip().startswith("http")]
            if not urls:
                await update.message.reply_text("Не нашёл ссылок. Отправь ссылки начинающиеся с http, каждую на новой строке. Или напиши *пропустить*.", parse_mode="Markdown")
                return
            await update.message.reply_text(f"🔍 Читаю {len(urls)} публикаци{'ю' if len(urls)==1 else 'и'}...")
            brief_state["client_profile"] = await _scrape_post_urls(urls)
        brief_state["step"] = "waiting_client_text"
        _brief_save(brief_state)
        await update.message.reply_text(
            "Шаг 2/3: Есть готовый текст от клиента? Отправь его сюда.\n\n"
            "Если нет — напиши *пропустить*.",
            parse_mode="Markdown"
        )
        return

    if brief_state.get("step") == "waiting_client_text":
        skip = user_text.lower() in ["пропустить", "пропусти", "skip", "-"]
        brief_state["client_text"] = "" if skip else user_text.strip()
        brief_state["step"] = "waiting_brief"
        _brief_save(brief_state)
        await update.message.reply_text(
            "Шаг 3/3: Отправь заполненный бриф — фото или PDF."
        )
        return

    if brief_state.get("step") in ["waiting_brief"] and user_text.lower() in ["стоп", "stop", "/brief stop"]:
        _brief_clear()
        await update.message.reply_text("Режим брифа выключен.")
        return

    # Режим toword — стоп
    if _toword_active() and user_text.lower() in ["стоп", "stop", "/toword stop"]:
        _toword_set(False)
        await update.message.reply_text("Режим конвертации выключен.")
        return

    # Ожидаем месяц для выписки
    pending = _statement_pending_load()
    if pending and pending.get("pdf_b64"):
        month_label = user_text.strip()
        _statement_pending_clear()
        await update.message.reply_text(f"📊 Анализирую выписку за {month_label}...")
        try:
            import asyncio
            from agents.statement_parser import analyze_with_claude, extract_expenses_structured
            parsed = {"pdf_base64": pending["pdf_b64"], "bank": pending.get("bank", "unknown")}

            def _do_analysis():
                analysis = analyze_with_claude(parsed, period=month_label)
                expenses = extract_expenses_structured(analysis, month_label)
                return analysis, expenses

            analysis, expenses = await asyncio.to_thread(_do_analysis)

            try:
                from agents.sheets_agent import append_expense_month
                append_expense_month(month_label, expenses.get("categories", {}),
                                     expenses.get("total_expenses", 0), expenses.get("total_income", 0))
                sheets_note = "✅ Данные сохранены в таблицу расходов."
            except Exception as se:
                logger.error(f"[Sheets] expense: {se}")
                sheets_note = "⚠️ Не удалось сохранить в таблицу."

            bank = parsed["bank"].upper() if parsed["bank"] != "unknown" else "банк"
            statement_analyses.append({"bank": bank, "period": month_label, "analysis": analysis})
            await send_long(update, f"📊 {bank} — {month_label}\n{'─'*30}\n\n{analysis}\n\n{sheets_note}")
        except Exception as e:
            logger.error(f"[Statement] error: {e}")
            await update.message.reply_text(f"Ошибка анализа: {e}")
        return

    # Режим аналитика
    if _analyst_active():
        if user_text.lower() in ["стоп", "stop", "/analyst stop"]:
            _analyst_set_active(False)
            await update.message.reply_text("🔴 Режим аналитика выключен.")
            return
        await _handle_analyst_message(update)
        return

    # Режим установки KPI
    ks = _load_kpiset()
    if ks["active"]:
        if user_text.lower() in ["отмена", "стоп", "cancel"]:
            _reset_kpiset()
            await update.message.reply_text("Отменено.")
            return
        from agents.sales_agent import set_kpi_plan, _parse_amount
        step = ks["step"]
        temp = ks["temp"]

        if step == "month":
            import re as _re3
            if not _re3.match(r"^\d{4}-\d{2}$", user_text.strip()):
                await update.message.reply_text("Напиши в формате: 2026-05")
                return
            temp["year_month"] = user_text.strip()
            _save_kpiset({"step": "grants_min", "temp": temp})
            await update.message.reply_text(
                "Grants KZ — нижний порог (0% бонус)?\n"
                "Пример: 1500000 или 1.5М"
            )
        elif step == "grants_min":
            temp["grants_min"] = _parse_amount(user_text)
            _save_kpiset({"step": "grants_max", "temp": temp})
            await update.message.reply_text(
                "Grants KZ — верхний порог (10% бонус)?\n"
                "Пример: 2000000 или 2М"
            )
        elif step == "grants_max":
            temp["grants_max"] = _parse_amount(user_text)
            _save_kpiset({"step": "tanda_min", "temp": temp})
            await update.message.reply_text(
                "Tanda Bilim — нижний порог (0% бонус)?\n"
                "Пример: 1500000 или 1.5М"
            )
        elif step == "tanda_min":
            temp["tanda_min"] = _parse_amount(user_text)
            _save_kpiset({"step": "tanda_max", "temp": temp})
            await update.message.reply_text(
                "Tanda Bilim — верхний порог (10% бонус)?\n"
                "Пример: 3000000 или 3М"
            )
        elif step == "tanda_max":
            temp["tanda_max"] = _parse_amount(user_text)
            ym = temp["year_month"]
            r1 = set_kpi_plan("grants_kz", ym, temp["grants_min"], temp["grants_max"])
            r2 = set_kpi_plan("tanda_bilim", ym, temp["tanda_min"], temp["tanda_max"])
            _reset_kpiset()
            await update.message.reply_text(f"✅ KPI план установлен на {ym}:\n\n{r1}\n{r2}")
        return

    # Режим добавления данных
    ad = _load_adddata()
    if ad["active"]:
        if user_text.lower() in ["отмена", "стоп", "cancel"]:
            _reset_adddata()
            await update.message.reply_text("Отменено.")
            return

        from agents.sales_agent import parse_month, save_historical_month, get_historical_raw, _parse_amount

        step = ad["step"]
        year = ad["year"]

        if step == "ask_month":
            month = parse_month(user_text)
            if not month:
                await update.message.reply_text("Не понял месяц. Напиши например: Май, Июнь, Январь")
                return
            _save_adddata({"month": month, "step": "grants_kz", "temp": {}})
            from agents.sales_agent import MONTH_NAMES
            await update.message.reply_text(
                f"📅 {MONTH_NAMES.get(month, month)} {year}\n\n"
                f"Grants KZ — выручка?\n"
                f"Пример: 1.5М или 800к или 0"
            )
            return

        elif step == "grants_kz":
            ad["temp"]["grants_kz"] = _parse_amount(user_text)
            _save_adddata({"step": "tanda_bilim", "temp": ad["temp"]})
            await update.message.reply_text("Tanda Bilim — выручка?\nПример: 1.2М или 0")
            return

        elif step == "tanda_bilim":
            ad["temp"]["tanda_bilim"] = _parse_amount(user_text)
            _save_adddata({"step": "ekonomist_media", "temp": ad["temp"]})
            await update.message.reply_text("Ekonomist Media — выручка?\nПример: 300к или 0")
            return

        elif step == "ekonomist_media":
            ad["temp"]["ekonomist_media"] = _parse_amount(user_text)
            try:
                save_historical_month(
                    year, ad["month"],
                    ad["temp"].get("grants_kz", 0),
                    ad["temp"].get("tanda_bilim", 0),
                    ad["temp"].get("ekonomist_media", 0),
                )
                from agents.sales_agent import push_historical_to_github
                push_historical_to_github()
                try:
                    from agents.sheets_agent import sync_historical_to_sheets
                    sync_historical_to_sheets()
                except Exception as se:
                    logger.error(f"[Sheets] historical sync error: {se}")
                    await update.message.reply_text(f"⚠️ Данные сохранены, но Google Sheets не обновился: {se}")
                _reset_adddata()
                updated = get_historical_raw(year)
                await update.message.reply_text(f"✅ Сохранено!\n\n{updated}")
            except Exception as e:
                await update.message.reply_text(f"❌ Ошибка: {e}")
            return


    # Режим создания счёта
    inv = _load_invoice()
    if inv["active"]:
        await _handle_invoice_dialog(update)
        return

    # Режим расчёта KPI
    kc = _load_kpicalc()
    if kc["active"]:
        if user_text.lower() in ["отмена", "стоп", "cancel"]:
            _reset_kpicalc()
            await update.message.reply_text("Отменено.")
            return
        from agents.sales_agent import calc_bonus, _parse_amount
        from agents.sales_agent import MONTH_NAMES
        step = kc["step"]
        temp = kc["temp"]

        if step == "ask_month":
            import re as _re_ym
            if user_text.lower() in ["текущий", "сейчас", "этот"]:
                ym = date.today().strftime("%Y-%m")
            elif _re_ym.match(r"^\d{4}-\d{2}$", user_text.strip()):
                ym = user_text.strip()
            else:
                await update.message.reply_text("Напиши в формате: 2026-04 или 'текущий'")
                return
            temp["year_month"] = ym
            _save_kpicalc({"step": "grants", "temp": temp})
            await update.message.reply_text(
                f"Grants KZ — фактическая выручка за {ym}?\n"
                f"Пример: 1.8М или 1800000"
            )
            return

        ym = temp["year_month"]

        if step == "grants":
            temp["grants_rev"] = _parse_amount(user_text)
            _save_kpicalc({"step": "tanda", "temp": temp})
            await update.message.reply_text(
                "Tanda Bilim — фактическая выручка за месяц?\n"
                "Пример: 2.5М или 0"
            )
        elif step == "tanda":
            temp["tanda_rev"] = _parse_amount(user_text)
            _reset_kpicalc()

            grants_rev = temp["grants_rev"]
            tanda_rev = temp["tanda_rev"]
            grants_bonus = calc_bonus("grants_kz", grants_rev, ym)
            tanda_bonus = calc_bonus("tanda_bilim", tanda_rev, ym)
            fixed = 200000 + 200000
            total = fixed + grants_bonus + tanda_bonus

            month_num = ym.split("-")[1]
            month_name = MONTH_NAMES.get(month_num, ym)

            lines = [f"💰 РАСЧЁТ БОНУСОВ — {month_name} {ym[:4]}", "─" * 32]
            lines.append(f"\n📌 Grants KZ")
            lines.append(f"  Выручка: {grants_rev:,.0f} ₸")
            lines.append(f"  Бонус: {grants_bonus:,.0f} ₸")
            lines.append(f"\n📌 Tanda Bilim")
            lines.append(f"  Выручка: {tanda_rev:,.0f} ₸")
            lines.append(f"  Бонус: {tanda_bonus:,.0f} ₸")
            lines.append(f"\n{'─' * 32}")
            lines.append(f"Фикс: {fixed:,.0f} ₸")
            lines.append(f"Бонусы: {grants_bonus + tanda_bonus:,.0f} ₸")
            lines.append(f"💵 К выплате: {total:,.0f} ₸")
            await update.message.reply_text("\n".join(lines))
            try:
                from agents.sheets_agent import upsert_kpi_row
                upsert_kpi_row(ym, grants_rev, tanda_rev,
                               grants_bonus, tanda_bonus, fixed, grants_bonus + tanda_bonus)
            except Exception as e:
                logger.error(f"[Sheets] KPI sync error: {e}")
        return

    # Запрос исторических продаж
    import re as _re
    hist_match = _re.search(r"20\d{2}", user_text)
    hist_keywords = ["продаж", "выручк", "отчет", "данные", "план", "статистик", "итог"]
    if any(w in user_text.lower() for w in ["два года", "2 года", "последних", "оба года"]):
        from agents.sales_agent import get_historical_two_years
        result = get_historical_two_years()
        await send_long(update, result)
        return
    if hist_match and any(w in user_text.lower() for w in hist_keywords):
        year = int(hist_match.group())
        from agents.sales_agent import get_historical_raw
        result = get_historical_raw(year)
        await send_long(update, result)
        return

    # Режим годового планирования
    yp = _load_yearplan()
    if yp["active"]:
        if any(w in user_text.lower() for w in ["отмена", "стоп", "cancel"]):
            _save_yearplan({"active": False, "data": [], "year": None, "waiting_year": False})
            await update.message.reply_text("Режим планирования отменён.")
            return

        # Ждём год
        if yp["waiting_year"]:
            import re as _re
            year_match = _re.search(r"20\d{2}", user_text)
            if not year_match:
                await update.message.reply_text("Напиши год цифрами, например: 2026")
                return
            yp["year"] = int(year_match.group())
            yp["waiting_year"] = False
            _save_yearplan(yp)
            await update.message.reply_text(
                f"✅ Составляем план на {yp['year']} год.\n\n"
                f"Теперь загрузи данные прошлых лет:\n"
                f"• PDF, Excel или CSV файл с выручкой\n"
                f"• Или напиши вручную:\n\n"
                f"2024: Grants KZ янв 800к, фев 1М...\n"
                f"Tanda Bilim: янв 600к, фев 700к...\n\n"
                f"Можно несколько файлов. Когда всё готово — напиши 'анализируй'."
            )
            return

        # Ждём данные или команду анализа
        trigger_words = ["анализируй", "готово", "составь план", "сделай план", "генерируй"]
        if any(w in user_text.lower() for w in trigger_words):
            if not yp["data"]:
                await update.message.reply_text("Сначала загрузи данные — файл или текст с выручкой прошлых лет.")
                return
            target_year = yp["year"]
            await update.message.reply_text(f"⏳ Анализирую данные и тренды рынка Казахстана для плана на {target_year}... Займёт минуту.")
            from agents.sales_agent import generate_annual_plan, save_annual_plan_as_targets
            combined = "\n\n---\n\n".join(yp["data"])
            plan = generate_annual_plan(combined, target_year)
            _save_yearplan({"active": False, "data": [], "year": None, "waiting_year": False})
            await send_long(update, plan)
            save_result = save_annual_plan_as_targets(plan, target_year)
            await update.message.reply_text(save_result)
            return
        else:
            yp["data"].append(user_text)
            _save_yearplan(yp)
            # Сохраняем в исторические данные — год ищем в тексте, fallback на "raw"
            import re as _re2
            year_in_text = _re2.search(r"20\d{2}", user_text)
            hist_year = int(year_in_text.group()) if year_in_text else (yp.get("year") or 0)
            if hist_year:
                from agents.sales_agent import save_historical_data
                save_historical_data(hist_year, user_text)
            chunks = len(yp["data"])
            await update.message.reply_text(
                f"✅ Данные добавлены ({chunks} блок(ов) загружено).\n"
                f"Загрузи ещё или напиши 'анализируй'."
            )
            return

    # Итог по нескольким выпискам
    if any(w in user_text.lower() for w in ["итог", "общая картина", "сводка", "суммируй"]):
        if not statement_analyses:
            await update.message.reply_text("Нет загруженных выписок.")
            return
        await send_combined_summary(update)
        return

    # Подтверждение / отмена
    if user_text.lower() in ["да", "ок", "ok", "yes", "подтверждаю", "верно"]:
        await update.message.reply_text("✅ Принято.")
        mark_responded()
        return
    if user_text.lower() in ["нет", "no", "отмена", "неверно"]:
        await update.message.reply_text("Окей, отменено. Отправьте правильные данные.")
        return

    # Проверяем — может это обновление балансов
    state = load_state()
    balance_keywords = ["каспи", "kaspi", "халык", "halyk", "фридом", "freedom"]
    if state.get("pending_update") or any(w in user_text.lower() for w in balance_keywords):
        await _process_balance_text(update, user_text)
        return

    # Всё остальное — оркестратор
    try:
        from core.orchestrator import run
        response = run(user_text, conversation_history)
        await send_long(update, response)
    except Exception as e:
        logger.error(f"Ошибка оркестратора: {e}")
        await update.message.reply_text(f"Ошибка: {e}")


# ─── PDF выписки ─────────────────────────────────────────────────────────────

async def handle_document(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    if not is_owner(update):
        return

    doc = update.message.document
    fname = doc.file_name.lower()

    # Режим брифа — PDF брифа
    if _brief_load().get("step") == "waiting_brief" and fname.endswith(".pdf"):
        tmp_path = None
        try:
            import base64
            file = await context.bot.get_file(doc.file_id)
            with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
                tmp_path = tmp.name
            await file.download_to_drive(tmp_path)
            with open(tmp_path, "rb") as f:
                pdf_b64 = base64.standard_b64encode(f.read()).decode()
            blocks = [{"type": "document", "source": {"type": "base64", "media_type": "application/pdf", "data": pdf_b64}}]
            await _process_brief_file(update, blocks)
        except Exception as e:
            await update.message.reply_text(f"❌ Ошибка: {e}")
        finally:
            if tmp_path:
                try: os.unlink(tmp_path)
                except: pass
        return

    # Режим toword — конвертируем PDF в Word
    if _toword_active() and fname.endswith(".pdf"):
        _toword_set(False)
        await update.message.reply_chat_action("upload_document")
        tmp_path = None
        try:
            file = await context.bot.get_file(doc.file_id)
            with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
                tmp_path = tmp.name
            await file.download_to_drive(tmp_path)
            with open(tmp_path, "rb") as f:
                pdf_bytes = f.read()
            await update.message.reply_text("🔍 Читаю текст из PDF...")
            text = await _extract_text_from_pdf_for_word(pdf_bytes)
            fname_out = doc.file_name.replace(".pdf", ".docx")
            word_buf = await _make_word_from_text(text, fname_out)
            await update.message.reply_document(document=word_buf, filename=fname_out, caption="✅ Word файл готов")
        except Exception as e:
            logger.error(f"[toword pdf] {e}")
            await update.message.reply_text(f"❌ Ошибка: {e}")
        finally:
            if tmp_path:
                try: os.unlink(tmp_path)
                except: pass
        return

    # Режим аналитика — читаем PDF и отправляем в диалог
    if _analyst_active() and fname.endswith(".pdf"):
        await update.message.reply_text(f"📄 Читаю PDF: {doc.file_name}...")
        try:
            import base64, anthropic as _ant, asyncio
            file = await context.bot.get_file(doc.file_id)
            with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
                tmp_path = tmp.name
            await file.download_to_drive(tmp_path)
            with open(tmp_path, "rb") as f:
                pdf_b64 = base64.standard_b64encode(f.read()).decode()
            os.unlink(tmp_path)

            caption = update.message.caption or "Проанализируй этот документ."

            history = _analyst_load_history()
            history.append({"role": "user", "content": [
                {"type": "document", "source": {"type": "base64", "media_type": "application/pdf", "data": pdf_b64}},
                {"type": "text", "text": caption}
            ]})

            await update.message.reply_chat_action("typing")

            def _call():
                client = _ant.Anthropic(api_key=os.getenv("ANTHROPIC_API_KEY"))
                biz_context = _build_analyst_context()
                system = ANALYST_SYSTEM + f"\n\n--- АКТУАЛЬНЫЕ ДАННЫЕ ---\n{biz_context}"
                resp = client.messages.create(
                    model="claude-opus-4-6",
                    max_tokens=2000,
                    system=system,
                    messages=history[-20:],
                )
                return resp.content[0].text

            answer = await asyncio.to_thread(_call)
            history.append({"role": "assistant", "content": answer})
            _analyst_save_history(history)
            await send_long(update, answer)
        except Exception as e:
            logger.error(f"[Analyst PDF] error: {e}")
            await update.message.reply_text(f"❌ Ошибка при чтении PDF: {e}")
        return

    # Режим годового планирования — принимаем любые файлы
    yp = _load_yearplan()
    if yp["active"]:
        await update.message.reply_text(f"📎 Получил файл: {doc.file_name}\nОбрабатываю...")
        tmp_path = None
        try:
            file = await context.bot.get_file(doc.file_id)
            suffix = "." + fname.split(".")[-1] if "." in fname else ".bin"
            with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
                tmp_path = tmp.name
            await file.download_to_drive(tmp_path)

            if fname.endswith(".pdf"):
                import base64
                with open(tmp_path, "rb") as f:
                    pdf_b64 = base64.standard_b64encode(f.read()).decode()
                # Извлекаем текст через Claude
                import anthropic as ant
                cl = ant.Anthropic(api_key=os.getenv("ANTHROPIC_API_KEY"))
                resp = cl.messages.create(
                    model="claude-opus-4-6", max_tokens=4096,
                    messages=[{"role": "user", "content": [
                        {"type": "document", "source": {"type": "base64", "media_type": "application/pdf", "data": pdf_b64}},
                        {"type": "text", "text": "Извлеки все данные о выручке и продажах из этого документа. Выведи в структурированном виде по месяцам и проектам."}
                    ]}]
                )
                extracted = resp.content[0].text
            elif fname.endswith((".csv", ".txt")):
                with open(tmp_path, "r", encoding="utf-8", errors="ignore") as f:
                    extracted = f.read()
            elif fname.endswith((".xlsx", ".xls")):
                try:
                    import openpyxl
                    wb = openpyxl.load_workbook(tmp_path, data_only=True)
                    rows = []
                    for sheet in wb.worksheets:
                        rows.append(f"[Лист: {sheet.title}]")
                        for row in sheet.iter_rows(values_only=True):
                            if any(cell is not None for cell in row):
                                rows.append("\t".join(str(c) if c is not None else "" for c in row))
                    extracted = "\n".join(rows)
                except ImportError:
                    extracted = f"Excel файл: {doc.file_name} (установи openpyxl для чтения)"
            else:
                extracted = f"Файл {doc.file_name} — не удалось прочитать автоматически."

            yp["data"].append(f"[Файл: {doc.file_name}]\n{extracted}")
            _save_yearplan(yp)
            if yp.get("year"):
                from agents.sales_agent import save_historical_data
                save_historical_data(yp["year"], f"[Файл: {doc.file_name}]\n{extracted}")
            await update.message.reply_text(
                f"✅ Файл обработан ({len(yp['data'])} блок(ов) загружено).\n"
                f"Загрузи ещё файлы или напиши данные вручную.\n"
                f"Когда всё готово — напиши 'анализируй'."
            )
        except Exception as e:
            logger.error(f"Ошибка обработки файла для плана: {e}")
            await update.message.reply_text(f"Ошибка обработки файла: {e}")
        finally:
            if tmp_path:
                try:
                    os.unlink(tmp_path)
                except Exception:
                    pass
        return

    if not fname.endswith(".pdf"):
        await update.message.reply_text("Поддерживаю только PDF выписки.")
        return

    caption = update.message.caption or ""

    # Если нет подписи — сохраняем PDF и спрашиваем месяц
    if not caption:
        await update.message.reply_text("📄 PDF получен. За какой месяц выписка? (например: Май 2026)")
        tmp_path = None
        try:
            file = await context.bot.get_file(doc.file_id)
            with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
                tmp_path = tmp.name
            await file.download_to_drive(tmp_path)
            from agents.statement_parser import parse_pdf
            parsed = parse_pdf(tmp_path)
            _statement_pending_save(parsed["pdf_base64"], parsed["bank"])
        except Exception as e:
            await update.message.reply_text(f"Ошибка: {e}")
        finally:
            if tmp_path:
                try: os.unlink(tmp_path)
                except: pass
        return

    count = len(statement_analyses) + 1
    await update.message.reply_text(f"📄 Выписка {count}: {doc.file_name}\nАнализирую...")

    tmp_path = None
    try:
        file = await context.bot.get_file(doc.file_id)
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
            tmp_path = tmp.name
        await file.download_to_drive(tmp_path)

        from agents.statement_parser import parse_pdf, analyze_with_claude, extract_expenses_structured

        parsed = parse_pdf(tmp_path)
        bank = parsed["bank"].upper() if parsed["bank"] != "unknown" else "банк"
        analysis = analyze_with_claude(parsed, period=caption)

        # Сохраняем структурированные данные в Google Sheets
        try:
            expenses = extract_expenses_structured(analysis, caption)
            from agents.sheets_agent import append_expense_month
            append_expense_month(caption, expenses.get("categories", {}),
                                 expenses.get("total_expenses", 0), expenses.get("total_income", 0))
        except Exception as se:
            logger.error(f"[Sheets] expense sync error: {se}")

        statement_analyses.append({
            "bank": bank,
            "period": caption,
            "analysis": analysis
        })

        header = f"📊 Выписка {count}: {bank} ({caption})"
        await send_long(update, f"{header}\n{'─'*30}\n\n{analysis}\n\n💡 Загружено: {len(statement_analyses)} шт. Данные сохранены в таблицу расходов.")

    except Exception as e:
        logger.error(f"Ошибка анализа: {e}")
        await update.message.reply_text(f"Ошибка: {e}")
    finally:
        if tmp_path:
            try:
                os.unlink(tmp_path)
            except Exception:
                pass


async def send_combined_summary(update: Update) -> None:
    import anthropic as ant
    client = ant.Anthropic(api_key=os.getenv("ANTHROPIC_API_KEY"))

    banks = ", ".join(set(a["bank"] for a in statement_analyses))
    all_analyses = "\n\n".join(
        f"--- Выписка {i+1}: {a['bank']} {a['period']} ---\n{a['analysis']}"
        for i, a in enumerate(statement_analyses)
    )

    prompt = f"""Объедини {len(statement_analyses)} выписки ({banks}) в единый отчёт:

{all_analyses}

Формат:
📅 [МЕСЯЦ ГОД]
  Доходы: +X ₸ | Расходы: -Y ₸ | Баланс: Z ₸
  Категории: Еда X ₸, Транспорт X ₸, ...

📊 ИТОГО: Доходы +X ₸, Расходы -Y ₸, Баланс Z ₸

💡 ВЫВОДЫ: 3 совета предпринимателю"""

    await update.message.reply_text(f"Объединяю {len(statement_analyses)} выписки...")

    resp = client.messages.create(
        model="claude-opus-4-6", max_tokens=4096,
        system="Финансовый советник. Объединяй данные точно, на русском.",
        messages=[{"role": "user", "content": prompt}]
    )

    statement_analyses.clear()
    await send_long(update, f"📊 ОБЩИЙ ОТЧЁТ ({banks})\n\n{resp.content[0].text}")


# ─── Запуск ──────────────────────────────────────────────────────────────────

async def on_startup(context):
    """Через 30 сек после старта проверяем пропущенные джобы за сегодня."""
    now = datetime.now(ALMATY_TZ)

    # После 23:00 — запрос балансов если ещё не запрашивали сегодня
    if now.hour >= 23:
        state = load_state()
        last_req = state.get("last_request", "")
        today_str = now.strftime("%Y-%m-%d")
        if today_str not in last_req:
            try:
                from agents.finance_agent import format_summary, _load
                data = _load()
                summary = format_summary(data)
                state["pending_update"] = True
                state["last_request"] = now.isoformat()
                save_state(state)
                text = (
                    f"🌙 Добрый вечер! Время обновить балансы.\n\n"
                    f"Текущее состояние:\n{summary}\n\n"
                    f"Отправьте новые балансы текстом или голосовым.\n"
                    f"Пример: каспи 150000, халык 800000"
                )
                await context.bot.send_message(chat_id=OWNER_ID, text=text)
                logger.info("Запрос балансов отправлен при старте (пропущен 23:00)")
            except Exception as e:
                logger.error(f"Не удалось запросить балансы: {e}")

    # После 17:00 — запрос продажнику если ещё не отправляли
    if now.hour >= 17 and SALESPERSON_ID != 0:
        from agents.sales_agent import SalesConversation, reset_daily_state, is_submitted_today
        conv = SalesConversation()
        if not conv.is_active() and not is_submitted_today():
            reset_daily_state()
            conv.reset()
            msg = conv.start()
            try:
                await context.bot.send_message(chat_id=SALESPERSON_ID, text=msg)
                logger.info("Запрос продажнику отправлен при старте (пропущен 17:00)")
            except Exception as e:
                logger.error(f"Не удалось отправить запрос продажнику: {e}")

    # После 19:00 — отчёт продаж владельцу если данные уже есть
    if now.hour >= 19:
        from agents.sales_agent import is_submitted_today, daily_report
        if is_submitted_today():
            try:
                report = daily_report()
                await context.bot.send_message(chat_id=OWNER_ID, text=f"📊 Дневной отчёт по продажам:\n\n{report}")
                logger.info("Отчёт продаж отправлен при старте (пропущен 19:00)")
            except Exception as e:
                logger.error(f"Не удалось отправить отчёт: {e}")


def main() -> None:
    token = os.getenv("TELEGRAM_BOT_TOKEN")
    if not token:
        raise ValueError("TELEGRAM_BOT_TOKEN не задан в .env")

    app = Application.builder().token(token).build()

    # Проверка пропущенных джобов через 30 сек после старта
    app.job_queue.run_once(on_startup, when=30, name="startup_check")

    # Курс USD/KZT: каждый день в 9:00
    app.job_queue.run_daily(
        auto_update_rate,
        time=time(9, 0, tzinfo=ALMATY_TZ),
        name="usd_rate_update"
    )

    # Балансы: 23:00 запрос, 12:00 повтор
    app.job_queue.run_daily(
        request_balances,
        time=time(23, 0, tzinfo=ALMATY_TZ),
        name="daily_balance_request"
    )
    app.job_queue.run_daily(
        reminder_balances,
        time=time(12, 0, tzinfo=ALMATY_TZ),
        name="balance_reminder"
    )

    # Продажи: 17:00 запрос у продажника
    app.job_queue.run_daily(
        request_sales_report,
        time=time(17, 0, tzinfo=ALMATY_TZ),
        name="sales_request"
    )
    # Каждые 30 мин: проверка зависшего диалога
    app.job_queue.run_repeating(
        check_stale_conversation,
        interval=1800,
        first=60,
        name="stale_conv_check"
    )

    # 19:00 отчёт владельцу или напоминание
    app.job_queue.run_daily(
        check_sales_19,
        time=time(19, 0, tzinfo=ALMATY_TZ),
        name="sales_check_19"
    )
    # 20:30 второе напоминание если нет данных
    app.job_queue.run_daily(
        remind_sales_2030,
        time=time(20, 30, tzinfo=ALMATY_TZ),
        name="sales_remind_2030"
    )
    # Еженедельный отчёт: понедельник 9:00
    app.job_queue.run_daily(
        weekly_sales_job,
        time=time(9, 0, tzinfo=ALMATY_TZ),
        days=(0,),
        name="weekly_sales"
    )
    # Ежемесячный отчёт: 1-е число 9:00
    app.job_queue.run_monthly(
        monthly_sales_job,
        when=time(9, 0, tzinfo=ALMATY_TZ),
        day=1,
        name="monthly_sales"
    )

    app.add_handler(CommandHandler("start", start))
    app.add_handler(CommandHandler("clear", cmd_clear))
    app.add_handler(CommandHandler("reset", cmd_reset))
    app.add_handler(CommandHandler("balances", cmd_balances))
    app.add_handler(CommandHandler("sales", cmd_sales))
    app.add_handler(CommandHandler("yearplan", cmd_yearplan))
    app.add_handler(CommandHandler("adddata", cmd_adddata))
    app.add_handler(CommandHandler("bonus", cmd_kpi))
    app.add_handler(CommandHandler("plan", cmd_setkpi))
    app.add_handler(CommandHandler("syncsheets", cmd_syncsheets))
    app.add_handler(CommandHandler("invoice", cmd_invoice))
    app.add_handler(CommandHandler("analyst", cmd_analyst))
    app.add_handler(CommandHandler("toword", cmd_toword))
    app.add_handler(CommandHandler("brief", cmd_brief))
    app.add_handler(CommandHandler("ceo", cmd_ceo))
    app.add_handler(CallbackQueryHandler(handle_ceo_callback, pattern="^ceo:"))
    app.add_handler(MessageHandler(filters.VOICE, handle_voice))
    app.add_handler(MessageHandler(filters.PHOTO, handle_photo))
    app.add_handler(MessageHandler(filters.Document.PDF, handle_document))
    app.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, handle_message))

    logger.info("Бот запущен. Планировщик: 23:00 и 12:00 по Алматы.")
    app.run_polling(allowed_updates=Update.ALL_TYPES)


if __name__ == "__main__":
    main()
